"""Формування звіту лабораторної роботи №2 за поточним ноутбуком."""

from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
NOTEBOOK_PATH = PROJECT_ROOT / "notebooks" / "euro_coin_classification.ipynb"
RESULTS_PATH = PROJECT_ROOT / "outputs" / "final_model_comparison.csv"
OUTPUT_PATH = PROJECT_ROOT / "lab2_report.docx"


MODEL_CELLS = [35, 37, 39, 41, 43, 45, 47, 49, 51, 53]
MODEL_NAMES = [
    "Логістична регресія",
    "Метод k-найближчих сусідів",
    "Метод опорних векторів",
    "Дерево рішень",
    "Випадковий ліс",
    "Надзвичайно рандомізовані дерева",
    "Градієнтний бустинг",
    "AdaBoost",
    "Гаусівський наївний Баєс",
    "Багатошаровий перцептрон",
]

MODEL_DESCRIPTIONS = [
    "Логістична регресія будує лінійні межі між класами. Вона використана як базовий інтерпретований метод для перевірки того, наскільки добре HOG-ознаки розділяються лінійною моделлю.",
    "KNN визначає клас монети за п'ятьма найближчими об'єктами навчальної вибірки. Метод залежить від відстаней між векторами, тому масштабування ознак є для нього особливо важливим.",
    "SVC із RBF-ядром будує нелінійні розділювальні межі з максимальним відступом між класами. Метод придатний для складної геометрії HOG-простору.",
    "Дерево рішень послідовно ділить простір ознак за окремими умовами. Модель легко інтерпретується, але може перенавчатися на невеликому наборі даних.",
    "Випадковий ліс об'єднує прогнози багатьох дерев, навчених на різних підвибірках об'єктів та ознак. Усереднення робить результат стабільнішим за одне дерево.",
    "Extra Trees також використовує ансамбль дерев, але сильніше рандомізує пороги поділу. Це зменшує схожість між деревами та часто забезпечує швидке навчання.",
    "Градієнтний бустинг будує дерева послідовно, причому кожне наступне дерево виправляє помилки попереднього ансамблю. Через послідовність метод має найбільший час навчання.",
    "AdaBoost підвищує вагу об'єктів, які раніше були класифіковані неправильно, і формує зважене голосування слабких класифікаторів.",
    "GaussianNB моделює розподіл кожної ознаки нормальним законом та припускає умовну незалежність ознак. Для пов'язаних HOG-компонент це припущення є спрощеним.",
    "MLPClassifier є невеликою нейронною мережею з одним прихованим шаром. Вона виконує нелінійне перетворення HOG-ознак без використання TensorFlow або PyTorch.",
]

MODEL_INTERPRETATIONS = [
    "Точність 0.8529 є найвищою серед усіх перевірених моделей. Це свідчить, що після масштабування класи у HOG-просторі значною мірою розділяються відносно простими межами. Для невеликої вибірки така модель узагальнила дані краще за складніші ансамблі, оскільки має меншу схильність підлаштовуватися під випадкові особливості окремих кропів.",
    "Точність 0.7353 показує, що локальна схожість HOG-векторів справді містить інформацію про номінал. Водночас часткове перекриття класів погіршує роботу правила найближчих сусідів: візуально подібна монета може опинитися ближче до об'єктів іншого класу, навіть після масштабування ознак.",
    "SVC досяг точності 0.7647 за короткого часу навчання 0.0463 с. Нелінійна межа допомогла врахувати складнішу геометрію класів, але повного розділення немає, оскільки частина монет має близькі контури й текстури. Саме поєднання достатньо високої точності та малого часу пояснює найкращий баланс цієї моделі за прийнятим у ноутбуці правилом.",
    "Точність 0.4265 є однією з найнижчих. Одне дерево приймає рішення через послідовні порогові поділи окремих HOG-компонентів, тоді як форма монети описується їх спільною конфігурацією. На невеликій вибірці це робить дерево нестійким і сприяє перенавчанню на деталях навчальних зображень.",
    "Випадковий ліс підвищив точність до 0.7500 порівняно з одним деревом. Результат підтверджує, що усереднення багатьох різних дерев зменшує вплив невдалих локальних поділів, хоча деревоподібне представлення все одно поступається найкращій лінійній моделі для цих HOG-ознак.",
    "Extra Trees отримав точність 0.7794, тобто найкращий результат серед ансамблів дерев у цьому експерименті. Додаткова випадковість порогів зменшила кореляцію між окремими деревами й допомогла краще контролювати дисперсію моделі на невеликому наборі.",
    "Градієнтний бустинг досяг точності 0.6471 і мав найбільший час навчання — 23.5517 с. На цьому наборі послідовне виправлення помилок не дало переваги: за малої кількості об'єктів наступні дерева можуть приділяти надмірну увагу складним або нетиповим кропам. Логістична регресія натомість використала загальну структуру стандартизованих HOG-ознак і тому узагальнила краще.",
    "AdaBoost показав найнижчу точність — 0.3235. Підвищення ваги раніше помилково класифікованих монет може концентрувати ансамбль на об'єктах із неоднозначною формою, освітленням або фоном. Якщо базові правила погано описують спільну структуру HOG-компонентів, послідовне підсилення не усуває цю проблему.",
    "GaussianNB навчався найшвидше — 0.0026 с, але його точність становила 0.6324. Швидкість пояснюється простим оцінюванням параметрів розподілів, а нижча точність — сильним припущенням про незалежність ознак, яке для сусідніх і пов'язаних HOG-компонентів виконується лише наближено.",
    "MLPClassifier отримав точність 0.7500 за часу навчання 0.0625 с. Нелінійне перетворення ознак виявилося корисним, однак 268 навчальних об'єктів недостатньо, щоб невелика нейронна мережа стабільно оцінювала велику кількість параметрів і перевершила регуляризовану логістичну регресію.",
]

CONFUSION_CAPTIONS = [
    "Матриця помилок логістичної регресії",
    "Матриця помилок методу k-найближчих сусідів",
    "Матриця помилок методу опорних векторів",
    "Матриця помилок дерева рішень",
    "Матриця помилок випадкового лісу",
    "Матриця помилок надзвичайно рандомізованих дерев",
    "Матриця помилок градієнтного бустингу",
    "Матриця помилок AdaBoost",
    "Матриця помилок гаусівського наївного Баєса",
    "Матриця помилок багатошарового перцептрона",
]

TSNE_MODEL_CAPTIONS = [
    "t-SNE прогнозованих класів логістичної регресії",
    "t-SNE прогнозованих класів методу k-найближчих сусідів",
    "t-SNE прогнозованих класів методу опорних векторів",
    "t-SNE прогнозованих класів дерева рішень",
    "t-SNE прогнозованих класів випадкового лісу",
    "t-SNE прогнозованих класів надзвичайно рандомізованих дерев",
    "t-SNE прогнозованих класів градієнтного бустингу",
    "t-SNE прогнозованих класів AdaBoost",
    "t-SNE прогнозованих класів гаусівського наївного Баєса",
    "t-SNE прогнозованих класів багатошарового перцептрона",
]


def set_run_font(run, name: str, size: float, bold: bool | None = None) -> None:
    """Встановлює однаковий шрифт для латиниці й кирилиці."""

    run.font.name = name
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def configure_document(document: Document) -> None:
    """Налаштовує сторінку та базові стилі відповідно до lab1_report."""

    section = document.sections[0]
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(0)
    section.footer_distance = Cm(0)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(0)

    for style_name, size in [("Heading 1", 14), ("Heading 2", 14), ("Heading 3", 14)]:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.keep_with_next = True


def add_body(document: Document, text: str, first_line: bool = True) -> None:
    """Додає основний абзац із вирівнюванням по ширині."""

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first_line:
        paragraph.paragraph_format.first_line_indent = Cm(1.25)
    run = paragraph.add_run(text)
    set_run_font(run, "Times New Roman", 14)


def add_heading(document: Document, text: str, level: int = 1, centered: bool = False) -> None:
    """Додає заголовок потрібного рівня."""

    paragraph = document.add_paragraph(style=f"Heading {level}")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    set_run_font(run, "Times New Roman", 14, bold=True)


def add_bullet(document: Document, text: str) -> None:
    """Додає маркований пункт у стилі попереднього звіту."""

    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Cm(0.75)
    paragraph.paragraph_format.first_line_indent = Cm(-0.4)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, "Times New Roman", 14)


def add_output(document: Document, number: int, title: str, content: str) -> None:
    """Додає нумерований текстовий вивід із моноширинним блоком."""

    title_paragraph = document.add_paragraph()
    title_paragraph.paragraph_format.keep_with_next = True
    title_run = title_paragraph.add_run(f"Вивід {number} — {title}")
    set_run_font(title_run, "Times New Roman", 14, bold=True)

    output_paragraph = document.add_paragraph()
    output_paragraph.paragraph_format.line_spacing = 1.0
    output_paragraph.paragraph_format.space_before = Pt(2)
    output_paragraph.paragraph_format.space_after = Pt(4)
    output_paragraph.paragraph_format.keep_together = True
    run = output_paragraph.add_run(content.strip())
    set_run_font(run, "Courier New", 8)


def set_cell_border(cell, color: str = "B7B7B7", size: str = "8") -> None:
    """Додає тонку сіру рамку навколо місця для скріншота."""

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def add_figure_placeholder(
    document: Document,
    number: int,
    caption: str,
    height_cm: float = 7.5,
) -> None:
    """Створює порожнє місце для ручного вставлення скріншота та підпис."""

    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Cm(16.5)
    row = table.rows[0]
    row.height = Cm(height_cm)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    cell = row.cells[0]
    cell.width = Cm(16.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_border(cell)
    placeholder = cell.paragraphs[0]
    placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = placeholder.add_run("МІСЦЕ ДЛЯ ВСТАВЛЕННЯ СКРІНШОТА")
    set_run_font(run, "Times New Roman", 11)
    run.font.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)

    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_before = Pt(3)
    caption_paragraph.paragraph_format.space_after = Pt(6)
    caption_run = caption_paragraph.add_run(f"Рисунок {number} — {caption}")
    set_run_font(caption_run, "Times New Roman", 14)


def stream_output(notebook: dict, cell_number: int) -> str:
    """Об'єднує всі текстові stream-output заданої комірки."""

    cell = notebook["cells"][cell_number - 1]
    parts = []
    for output in cell.get("outputs", []):
        if output.get("output_type") == "stream":
            parts.append("".join(output.get("text", [])))
    return "\n".join(part.strip() for part in parts if part.strip())


def plain_output(notebook: dict, cell_number: int) -> list[str]:
    """Повертає текстові представлення display_data та execute_result."""

    values = []
    cell = notebook["cells"][cell_number - 1]
    for output in cell.get("outputs", []):
        if output.get("output_type") not in {"display_data", "execute_result"}:
            continue
        value = output.get("data", {}).get("text/plain")
        if value is None:
            continue
        values.append("".join(value) if isinstance(value, list) else str(value))
    return values


def model_output(notebook: dict, cell_number: int) -> str:
    """Формує блок часу, accuracy та classification report однієї моделі."""

    stream = stream_output(notebook, cell_number)
    reports = [value for value in plain_output(notebook, cell_number) if "точність" in value]
    report = reports[0] if reports else ""
    stream_lines = [
        line
        for line in stream.splitlines()
        if line.startswith("Час навчання") or line.startswith("Точність на тестовій")
    ]
    return "\n".join(stream_lines + ([report.strip()] if report else []))


def describe_excerpt(notebook: dict) -> str:
    """Вирізає читабельний фрагмент реального describe-output."""

    values = plain_output(notebook, 30)
    table_text = next(value for value in values if "HOG_0" in value)
    lines = table_text.splitlines()
    excerpt = []
    for line in lines:
        if line.startswith("          HOG_7"):
            break
        excerpt.append(line)
    excerpt.append("...")
    excerpt.append("[8 рядків x 1764 стовпці]")
    return "\n".join(excerpt)


def final_results_text(results: list[dict[str, str]]) -> str:
    """Форматує фінальну таблицю accuracy та часу як у текстовому output."""

    rows = [f"{'Модель':<38} {'Точність':>10} {'Час, с':>10}"]
    rows.append("-" * 62)
    for row in results:
        rows.append(
            f"{row['model name']:<38} "
            f"{float(row['test accuracy']):>10.4f} "
            f"{float(row['training time']):>10.4f}"
        )
    return "\n".join(rows)


def add_title_page(document: Document) -> None:
    """Створює титульну сторінку за зразком lab1_report."""

    for text in [
        "Міністерство освіти і науки України",
        "Львівський національний університет ім. І. Франка",
        "Факультет електроніки та комп’ютерних технологій",
        "Кафедра системного проектування",
    ]:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        set_run_font(run, "Times New Roman", 14)

    for _ in range(5):
        document.add_paragraph()

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = paragraph.add_run("ЗВІТ")
    set_run_font(run, "Times New Roman", 14, bold=True)
    run.add_break()
    run = paragraph.add_run("Про виконання лабораторної роботи №2")
    set_run_font(run, "Times New Roman", 14)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = paragraph.add_run(
        "“Дослідження різних методів багатокласової класифікації "
        "для визначення номіналу монети євро на основі набору даних EURO coins”"
    )
    set_run_font(run, "Times New Roman", 14)

    for _ in range(5):
        document.add_paragraph()

    for text, bold in [
        ("Виконав:", True),
        ("студент групи ФЕП-32с", False),
        ("Коваль С. А.", False),
        ("", False),
        ("Перевірив:", True),
        ("асистент Парубочий В. О.", False),
    ]:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run(text)
        set_run_font(run, "Times New Roman", 14, bold=bold)

    for _ in range(1):
        document.add_paragraph()

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("ЛЬВІВ – 2026")
    set_run_font(run, "Times New Roman", 14, bold=True)
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def build_report() -> None:
    """Збирає завершений DOCX без вставлених графічних зображень."""

    notebook = json.loads(NOTEBOOK_PATH.read_text(encoding="utf-8"))
    with RESULTS_PATH.open(encoding="utf-8", newline="") as csv_file:
        results = list(csv.DictReader(csv_file))
    document = Document()
    configure_document(document)
    add_title_page(document)

    add_heading(document, "Мета роботи", level=1)
    add_body(
        document,
        "Дослідити процес перетворення набору зображень із YOLO-анотаціями "
        "на задачу багатокласової класифікації, виконати виділення та аналіз "
        "HOG-ознак, підготувати дані, навчити десять методів класифікації й "
        "порівняти їх за точністю та часом навчання.",
    )

    add_heading(document, "Середовище розробки", level=1)
    add_bullet(document, "Мова програмування: Python")
    add_bullet(document, "Середовище: Jupyter Notebook")
    add_bullet(
        document,
        "Бібліотеки: NumPy, pandas, Pillow, Matplotlib, seaborn, scikit-image, scikit-learn",
    )
    add_bullet(document, "Набір даних: EURO coins із YOLO-анотаціями")

    add_heading(document, "Короткий опис набору даних", level=1)
    add_body(
        document,
        "Набір містить 150 фотографій монет євро та 150 відповідних файлів "
        "розмітки. У кожному рядку YOLO-файлу задано class_id, нормалізовані "
        "координати центра рамки, її ширину та висоту. На одній фотографії може "
        "бути кілька монет, тому для задачі класифікації кожен анотований об'єкт "
        "було вирізано в окремий кроп розміром 64x64 пікселі.",
    )
    for item in [
        "0 — 1 цент",
        "1 — 2 центи",
        "2 — 5 центів",
        "3 — 10 центів",
        "4 — 20 центів",
        "5 — 50 центів",
        "6 — 1 євро",
        "7 — 2 євро",
    ]:
        add_bullet(document, item)

    document.add_page_break()
    add_heading(document, "Хід роботи", level=1, centered=True)

    # 1. Аналіз набору даних
    add_heading(document, "1. Аналіз набору даних", level=1)
    add_body(
        document,
        "Початковий аналіз потрібний, щоб переконатися, що майбутня модель "
        "навчатиметься саме на монетах, а не на випадкових ділянках фону. "
        "Очікуваним результатом цього етапу є коректний кроп для кожної "
        "валідної YOLO-анотації та відсутність порожніх зображень.",
    )
    add_body(
        document,
        "Для кожного зображення було знайдено однойменний файл розмітки. "
        "Нормалізовані координати YOLO переведено у піксельні координати з "
        "урахуванням фактичної ширини та висоти фотографії. Рамки перевірено "
        "на коректність, після чого монети вирізано та масштабовано до 64x64.",
    )
    add_output(document, 1, "Приклад сформованого кропу", stream_output(notebook, 5))
    add_body(
        document,
        "Перший сформований об'єкт має очікуваний розмір 64x64 і зберігає "
        "зв'язок із класом, вихідним файлом та піксельною рамкою. Отже, "
        "перетворення анотації не втрачає інформацію, потрібну для подальшої "
        "перевірки походження кожного прикладу.",
    )
    add_output(
        document,
        2,
        "Підсумкова перевірка та розподіл об'єктів за класами",
        stream_output(notebook, 7),
    )
    add_body(
        document,
        "У результаті зі 150 зображень отримано 336 об'єктів восьми класів, "
        "а всі лічильники помилок дорівнюють нулю. Це означає, що технічні "
        "проблеми читання даних не вплинули на склад вибірки, тому подальший "
        "аналіз відображає властивості самих зображень.",
    )
    add_figure_placeholder(document, 1, "Розподіл вирізаних монет за номіналами", 7.5)
    add_body(
        document,
        "Найбільше об'єктів належить класу 5 центів — 56 кропів, а найменше "
        "класу 50 центів — 27 кропів. Такий дисбаланс означає, що модель бачить "
        "понад удвічі більше прикладів найчисельнішого класу й може краще "
        "пристосуватися саме до нього, тоді як оцінки для найменшого класу є "
        "менш стабільними. Саме тому при поділі даних надалі збережено пропорції "
        "класів за допомогою stratify.",
    )
    add_body(
        document,
        "Після кількісної перевірки необхідно візуально оцінити самі кропи. "
        "Це дає змогу побачити помилки рамок, які могли б не проявитися у "
        "лічильниках, але спотворили б ознаки та навчання моделей.",
    )
    add_figure_placeholder(document, 2, "Приклади кропів монет для кожного класу", 13.0)
    add_body(
        document,
        "Візуальна перевірка показує, що рамки відповідають монетам, а кропи "
        "містять об'єкти потрібного номіналу без порожніх або пошкоджених зображень.",
    )

    # 2. Підготовка даних
    add_heading(document, "2. Підготовка даних", level=1)
    add_heading(document, "2.1 Перетворення зображень у числові ознаки", level=2)
    add_body(
        document,
        "Класифікатор не може безпосередньо працювати із зображенням як із "
        "візуальним об'єктом, тому його потрібно подати як числовий вектор. "
        "Перехід до відтінків сірого зменшує вплив кольору й залишає структуру "
        "яскравості, за якою можна описувати форму та рельєф монети.",
    )
    add_output(document, 3, "Результат перетворення у відтінки сірого", stream_output(notebook, 17))
    add_body(
        document,
        "Матриця зберігає 336 об'єктів розміром 64x64, а значення лежать у "
        "вказаному у виводі діапазоні. Отже, всі кропи приведено до єдиного "
        "числового формату без зміни кількості навчальних прикладів.",
    )
    add_body(
        document,
        "Найпростіший спосіб подати зображення моделі — послідовно розгорнути "
        "всі його пікселі. Таке представлення використовується як зрозуміла "
        "базова перевірка, хоча воно чутливе до освітлення та невеликих зсувів.",
    )
    add_output(document, 4, "Розмір матриці розгорнутих пікселів", stream_output(notebook, 19))
    add_body(
        document,
        "Кожен кроп описується 4096 значеннями яскравості. Така кількість "
        "параметрів значно перевищує кількість об'єктів, а сусідні пікселі "
        "сильно залежать один від одного, тому raw-pixel представлення не було "
        "обрано основним для порівняння класичних моделей.",
    )
    add_body(
        document,
        "Для монет важливими є круглий контур, цифри та рельєфні елементи, "
        "тому доцільно описувати не окремі яскравості, а напрямки локальних "
        "градієнтів. Від HOG очікуємо більшої стійкості до помірних змін "
        "освітлення та кращого відображення геометрії номіналу.",
    )
    add_figure_placeholder(document, 3, "Монета у відтінках сірого та візуалізація HOG-ознак", 7.0)
    add_body(
        document,
        "Візуалізація HOG підкреслює напрямки країв замість точного відтворення "
        "текстури. Саме така абстракція корисна для порівняння монет, бо вона "
        "зосереджує модель на контурах написів і рельєфу.",
    )
    hog_output = stream_output(notebook, 21)
    add_output(document, 5, "Перевірка матриці HOG-ознак і цільового вектора", hog_output + "\n" + stream_output(notebook, 23))
    add_body(
        document,
        "Кожна монета представлена 1764 HOG-компонентами, а матриця X і вектор "
        "y містять однакові 336 об'єктів. Відсутність пропущених значень і "
        "повністю однакових рядків дає змогу переходити до поділу вибірки без "
        "додаткового відновлення або вилучення спостережень.",
    )

    add_heading(document, "2.2 Кодування, масштабування та поділ вибірки", level=2)
    add_body(
        document,
        "Перед навчанням текстові назви класів потрібно подати у формі, яку "
        "підтримують алгоритми scikit-learn, не створюючи при цьому штучного "
        "порядку серед вхідних ознак. Тому LabelEncoder застосовується лише до "
        "цільової змінної, а числові HOG-компоненти не потребують кодування.",
    )
    add_body(
        document,
        "Цільові мітки закодовано за допомогою LabelEncoder. Вхідні HOG-ознаки "
        "є числовими, тому кодування категоріальних вхідних властивостей не "
        "проводилось.",
    )
    add_body(
        document,
        "Компоненти HOG мають різні діапазони, тому без масштабування ознаки з "
        "більшою амплітудою непропорційно впливали б на відстані KNN, межу SVC, "
        "коефіцієнти логістичної регресії та градієнтне навчання MLP. "
        "StandardScaler приводить їх до зіставного масштабу, щоб алгоритми "
        "реагували на інформацію ознаки, а не на одиниці її вимірювання.",
    )
    add_body(
        document,
        "Окрему тестову частину формуємо для перевірки узагальнення на даних, "
        "яких модель не бачила під час fit. Через нерівномірний розподіл "
        "номіналів застосовано стратифікований поділ 80/20, щоб обидві частини "
        "зберігали початкові пропорції класів.",
    )
    add_output(document, 6, "Кодування класів і розміри навчальної та тестової вибірок", stream_output(notebook, 25))
    add_body(
        document,
        "Після кодування кожному з восьми номіналів відповідає однозначна "
        "числова мітка. Навчальна вибірка містить 268 об'єктів, а тестова — 68; "
        "це залишає більшу частину даних для оцінювання параметрів і водночас "
        "забезпечує незалежну перевірку всіх моделей.",
    )

    add_heading(document, "2.3 Видалення неінформативних ознак", level=2)
    add_body(
        document,
        "Ознака, яка має однакове значення для всіх навчальних об'єктів, не може "
        "допомогти розрізнити класи, але збільшує розмірність обчислень. Тому "
        "перед навчанням перевіряємо HOG-компоненти на нульову дисперсію, "
        "очікуючи вилучити лише повністю сталі стовпці.",
    )
    add_body(
        document,
        "До навчальної частини застосовано VarianceThreshold із порогом 0.0. "
        "Фільтр мав видалити компоненти, значення яких є сталими для всіх "
        "навчальних об'єктів.",
    )
    add_output(document, 7, "Результат застосування VarianceThreshold", stream_output(notebook, 28))
    add_body(
        document,
        "Кількість ознак до і після фільтрації залишилася рівною 1764, тобто "
        "жодної сталої компоненти не виявлено. Це не є недоліком етапу: "
        "результат підтверджує, що кожна HOG-ознака хоча б частково змінюється "
        "між навчальними монетами й автоматичне скорочення не потрібне.",
    )

    add_heading(document, "2.4 Статистичний аналіз HOG-ознак", level=2)
    add_body(
        document,
        "Статистичний опис потрібний, щоб перевірити фактичний масштаб, розкид і "
        "крайні значення після стандартизації. Очікуємо середні значення поблизу "
        "нуля та порівнювані стандартні відхилення, хоча окремі компоненти можуть "
        "мати асиметрію через рідкісні сильні контури.",
    )
    add_body(
        document,
        "Для нормалізованої навчальної матриці сформовано статистичний опис "
        "методом describe(). Таблиця містить кількість, середнє, стандартне "
        "відхилення, мінімум, квартилі та максимум для 1764 компонентів. "
        "Нижче наведено фрагмент фактичного виводу.",
    )
    add_output(document, 8, "Статистичний опис HOG-ознак (фрагмент)", describe_excerpt(notebook))
    add_body(
        document,
        "У наведеному фрагменті середні значення близькі до нуля, а стандартні "
        "відхилення — до одиниці, що узгоджується з роботою StandardScaler. "
        "Помітні максимальні значення показують, що для окремих монет деякі "
        "напрямки градієнтів виражені значно сильніше, ніж у більшості вибірки.",
    )
    add_body(
        document,
        "Окремо порівнюємо середню яскравість кропів між номіналами. Це не "
        "основна навчальна ознака, але вона допомагає оцінити, наскільки умови "
        "зйомки й матеріал монети можуть створювати внутрішньокласовий розкид.",
    )
    add_figure_placeholder(document, 4, "Розподіл середньої яскравості пікселів монет за номіналами", 7.5)
    add_body(
        document,
        "Boxplot показує відмінності середньої яскравості між класами та розкид "
        "всередині класів. Ця характеристика описує grayscale-пікселі, а не HOG, "
        "і відображає вплив освітлення, поверхні монети та фону.",
    )

    # 3. Навчання моделей
    add_heading(document, "3. Навчання 10 методів класифікації", level=1)
    add_body(
        document,
        "Порівняння різних алгоритмів на однакових ознаках дає змогу відокремити "
        "вплив моделі від впливу підготовки даних. Очікуємо з'ясувати, чи достатньо "
        "для HOG-векторів простої лінійної межі, чи перевагу матимуть нелінійні "
        "методи та ансамблі дерев.",
    )
    add_body(
        document,
        "Усі моделі навчались на однаковій матриці HOG-ознак і тому самому "
        "стратифікованому поділі. Для оцінювання використано accuracy, precision, "
        "recall, F1-міру та матрицю помилок. Час вимірювався лише для операції fit.",
    )

    output_number = 9
    figure_number = 5
    for index, (cell_number, model_name, description, interpretation, caption) in enumerate(
        zip(
            MODEL_CELLS,
            MODEL_NAMES,
            MODEL_DESCRIPTIONS,
            MODEL_INTERPRETATIONS,
            CONFUSION_CAPTIONS,
        ),
        start=1,
    ):
        add_heading(document, f"3.{index} {model_name}", level=2)
        add_body(document, description)
        add_output(
            document,
            output_number,
            f"Результати моделі «{model_name}»",
            model_output(notebook, cell_number),
        )
        add_figure_placeholder(document, figure_number, caption, 8.5)
        add_body(document, interpretation)
        output_number += 1
        figure_number += 1

    # 4. Візуалізація
    add_heading(document, "4. Візуалізація результатів", level=1)
    add_heading(document, "4.1 Порівняння точності та часу навчання", level=2)
    add_body(
        document,
        "Окремі classification report дають детальну оцінку кожної моделі, але "
        "для підсумкового вибору потрібне спільне представлення. Таблиця й "
        "діаграми мають показати не лише лідера за точністю, а й ціну цього "
        "результату в часі навчання.",
    )
    add_body(
        document,
        "Фінальну таблицю відсортовано за тестовою точністю у спадному порядку. "
        "Час навчання залежить від обладнання, але в межах одного запуску дає "
        "можливість порівняти обчислювальну складність методів.",
    )
    add_output(document, 19, "Фінальна таблиця точності та часу навчання", final_results_text(results))
    add_body(
        document,
        "Таблиця підтверджує, що логістична регресія посіла перше місце з "
        "точністю 0.8529, тоді як Extra Trees і SVC сформували наступну групу "
        "сильних результатів. Найменший час має GaussianNB, але його точність "
        "нижча; тому найшвидший метод не є автоматично найкращим для практичної "
        "класифікації.",
    )
    add_figure_placeholder(
        document,
        figure_number,
        "Порівняння точності моделей на тестовій вибірці",
        8.0,
    )
    figure_number += 1
    add_body(
        document,
        "Діаграма точності робить розрив між методами наочним: простіша "
        "логістична регресія випередила складні ансамблі, а дерево рішень і "
        "AdaBoost залишилися внизу рейтингу. Це підтримує висновок, що для "
        "невеликої високорозмірної HOG-вибірки складність моделі сама по собі "
        "не гарантує кращого узагальнення.",
    )
    add_figure_placeholder(
        document,
        figure_number,
        "Порівняння часу навчання моделей",
        8.0,
    )
    figure_number += 1
    add_body(
        document,
        "За часом особливо виділяється градієнтний бустинг із 23.5517 с, хоча "
        "його точність становить лише 0.6471. SVC навчився за 0.0463 с і досяг "
        "0.7647, тому в межах цього експерименту краще поєднує якість і "
        "обчислювальні витрати.",
    )

    add_heading(document, "4.2 t-SNE аналіз", level=2)
    add_body(
        document,
        "Числова точність не показує, як саме об'єкти розташовані в просторі "
        "ознак. t-SNE потрібний, щоб зменшити 1764 виміри до двох і візуально "
        "перевірити, чи утворюють номінали окремі групи та де виникають зони "
        "неоднозначності.",
    )
    add_body(
        document,
        "Метод t-SNE зменшує 1764 HOG-ознаки до двох координат, намагаючись "
        "розташувати схожі об'єкти поруч. Perplexity автоматично обмежено "
        "розміром тестової вибірки. Компактні одноколірні групи вказують на "
        "добре відокремлені класи, а перекриття кольорів — на подібність монет "
        "і можливі помилки класифікації.",
    )
    add_output(document, 20, "Параметри та форма t-SNE представлення", stream_output(notebook, 64))
    add_body(
        document,
        "Для 68 тестових об'єктів автоматично обрано безпечне значення "
        "perplexity 22, після чого отримано двовимірне представлення форми "
        "(68, 2). Отже, параметр узгоджений із розміром тестової вибірки, а "
        "кожна монета має координати для подальшого порівняння класів.",
    )
    add_figure_placeholder(
        document,
        figure_number,
        "t-SNE тестових монет за справжніми класами",
        8.5,
    )
    figure_number += 1
    add_body(
        document,
        "На графіку справжніх класів поряд із локальними групами видно "
        "перекриття кольорів. Це означає, що частина різних номіналів має "
        "подібні HOG-описи: спільна кругла форма, близькі контури рельєфу, "
        "повороти та умови освітлення не дають повністю розділити всі монети.",
    )
    add_body(
        document,
        "Тепер ті самі координати розфарбовуються прогнозами кожної моделі. "
        "Порівняння зі справжніми кольорами показує, чи зберігає класифікатор "
        "локальну структуру груп, чи об'єднує неоднозначні області під однією "
        "міткою.",
    )
    for caption in TSNE_MODEL_CAPTIONS:
        add_figure_placeholder(document, figure_number, caption, 8.0)
        figure_number += 1

    add_body(
        document,
        "На t-SNE графіках класи частково перекриваються. Саме в таких змішаних "
        "ділянках матриці помилок фіксують взаємні хибні прогнози: це два "
        "представлення тієї самої проблеми подібності HOG-векторів. Через "
        "спільну круглу форму, близькі напрямки країв, різне освітлення, "
        "повороти, фон і сторону монети візуально схожі об'єкти різних класів "
        "опиняються поруч. Графіки прогнозів найточніших моделей краще "
        "відтворюють структуру справжніх класів, тоді як слабші моделі частіше "
        "зафарбовують змішані області однією неправильною міткою.",
    )

    document.add_page_break()
    add_heading(document, "Висновок", level=1)
    add_body(
        document,
        "Під час виконання лабораторної роботи набір EURO coins було успішно "
        "перетворено із задачі детекції об'єктів на задачу багатокласової "
        "класифікації. На основі YOLO-анотацій сформовано 336 кропів восьми "
        "номіналів. Зображення переведено у відтінки сірого, сформовано raw-pixel "
        "та HOG-представлення, а для навчання використано 1764 HOG-ознаки.",
    )
    add_body(
        document,
        "Після масштабування та стратифікованого поділу було навчено десять "
        "методів. Найкращу тестову точність показала логістична регресія — "
        "0.8529. Найменший час fit у поточному запуску мав GaussianNB — 0.0026 с, "
        "а найкращий баланс точності й часу за прийнятим правилом отримав SVC. "
        "Extra Trees, SVC, Random Forest і MLP також показали достатньо високі "
        "результати, тоді як окреме дерево та AdaBoost виявилися менш стійкими.",
    )
    add_body(
        document,
        "Перевага логістичної регресії над градієнтним бустингом показує, що на "
        "невеликій високорозмірній вибірці стандартизовані HOG-ознаки краще "
        "узагальнюються відносно простою регуляризованою межею. Дерево рішень "
        "виявилося чутливим до окремих порогових поділів, а AdaBoost — до "
        "складних і неоднозначних прикладів, вагу яких він послідовно збільшує.",
    )
    add_body(
        document,
        "Основними обмеженнями роботи є невеликий розмір і дисбаланс набору, "
        "різне освітлення, положення, повороти, фони та сторони монет. Через це "
        "візуально подібні номінали частково перекриваються у t-SNE просторі й "
        "утворюють відповідні помилки в матрицях класифікації.",
    )
    add_body(
        document,
        "Подальше дослідження можна покращити збільшенням кількості розмічених "
        "монет і застосуванням аугментації поворотів, масштабу та освітлення. "
        "CNN могла б автоматично навчати складніші просторові й текстурні "
        "ознаки замість фіксованого HOG-опису. Водночас HOG і класичні алгоритми "
        "потребують менше даних та обчислень, залишаються прозорішими для "
        "порівняння і тому добре відповідають меті цієї університетської роботи.",
    )

    document.save(OUTPUT_PATH)
    print(f"Створено: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_report()
